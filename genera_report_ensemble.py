"""
Script per generare un report Word dettagliato sulle previsioni ensemble.
Spiega teoria, formule, percentili e analisi dei dati ICON-EPS per Londra.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import numpy as np

# ── Dati ensemble (40 members, ICON Seamless EPS) ──────────────────────────

data = {
    "2026-03-14": [
        11.2, 11.0, 10.9, 11.1, 10.6, 11.4, 9.7, 11.4, 10.4, 10.5,
        10.6, 10.7, 10.7, 10.9, 10.6, 10.4, 10.6, 11.2, 10.9, 9.8,
        10.9, 10.1, 11.0, 11.1, 11.1, 9.9, 11.1, 10.4, 10.3, 10.5,
        10.4, 10.5, 11.1, 10.8, 10.0, 10.3, 10.7, 10.5, 11.0, 10.6,
    ],
    "2026-03-15": [
        11.1, 10.4, 11.6, 9.4, 10.6, 10.8, 10.8, 11.3, 10.0, 10.5,
        10.2, 10.4, 11.0, 10.0, 10.4, 10.3, 10.3, 10.6, 11.2, 12.2,
        10.6, 10.6, 10.8, 10.3, 11.0, 10.4, 10.1, 11.0, 12.1, 10.5,
        11.1, 10.7, 10.6, 11.8, 11.2, 10.2, 10.5, 10.2, 11.1, 9.5,
    ],
    "2026-03-16": [
        10.7, 10.1, 11.5, 11.0, 11.6, 11.1, 10.2, 11.1, 11.0, 10.6,
        11.0, 11.6, 10.8, 10.7, 10.5, 9.7, 10.5, 11.1, 10.7, 11.1,
        10.7, 11.4, 12.4, 11.1, 10.8, 9.3, 12.5, 10.7, 10.9, 11.0,
        9.8, 9.9, 11.3, 12.2, 11.6, 10.3, 11.0, 10.0, 10.3, 11.0,
    ],
}


# ── Funzioni di utilità ─────────────────────────────────────────────────────

def set_cell_shading(cell, color_hex):
    """Applica colore di sfondo a una cella."""
    shading = cell._element.get_or_add_tcPr()
    shading_elm = shading.makeelement(
        qn("w:shd"),
        {qn("w:fill"): color_hex, qn("w:val"): "clear"},
    )
    shading.append(shading_elm)


def style_header_row(row, color_hex="1F4E79"):
    """Formatta riga di intestazione tabella."""
    for cell in row.cells:
        set_cell_shading(cell, color_hex)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(9)


def add_table_row(table, values, bold=False, center=True):
    """Aggiunge riga a una tabella."""
    row = table.add_row()
    for i, val in enumerate(values):
        cell = row.cells[i]
        p = cell.paragraphs[0]
        if center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(val))
        run.font.size = Pt(9)
        if bold:
            run.font.bold = True
    return row


def compute_stats(values):
    """Calcola tutte le statistiche per un array di valori ensemble."""
    arr = np.array(values)
    return {
        "media": np.mean(arr),
        "mediana": np.median(arr),
        "std": np.std(arr, ddof=1),
        "min": np.min(arr),
        "max": np.max(arr),
        "range": np.max(arr) - np.min(arr),
        "p10": np.percentile(arr, 10),
        "p25": np.percentile(arr, 25),
        "p75": np.percentile(arr, 75),
        "p90": np.percentile(arr, 90),
        "n_above_11": int(np.sum(arr >= 11.0)),
        "n_10_11": int(np.sum((arr >= 10.0) & (arr < 11.0))),
        "n_below_10": int(np.sum(arr < 10.0)),
        "n_above_12": int(np.sum(arr >= 12.0)),
    }


# ── Creazione documento ─────────────────────────────────────────────────────

doc = Document()

# Margini pagina
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

# ═══════════════════════════════════════════════════════════════════════════
# TITOLO
# ═══════════════════════════════════════════════════════════════════════════

title = doc.add_heading("Previsioni Ensemble e Calcolo delle Probabilita", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run("Analisi della temperatura massima a Londra (51.51N, 0.06E)")
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = sub2.add_run("Modello: ICON Seamless EPS  |  Fonte: Open-Meteo Ensemble API  |  14-16 Marzo 2026")
run2.font.size = Pt(10)
run2.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_paragraph("")  # spaziatura

# ═══════════════════════════════════════════════════════════════════════════
# SEZIONE 1 – TEORIA
# ═══════════════════════════════════════════════════════════════════════════

doc.add_heading("1. Cos'e una previsione ensemble?", level=1)

doc.add_paragraph(
    "Una previsione ensemble e un metodo probabilistico utilizzato in meteorologia "
    "per quantificare l'incertezza delle previsioni del tempo. Invece di eseguire "
    "una singola simulazione del modello atmosferico, si eseguono N simulazioni "
    "parallele (chiamate \"members\" o \"membri\"), ciascuna con condizioni iniziali "
    "leggermente perturbate."
)

doc.add_paragraph(
    "L'idea fondamentale e che lo stato dell'atmosfera non e mai conosciuto con "
    "precisione perfetta: le osservazioni (satelliti, radiosonde, stazioni meteo) "
    "contengono sempre errori di misura. Modificando leggermente queste condizioni "
    "iniziali entro il margine di errore plausibile, si ottengono traiettorie "
    "meteorologiche diverse che rappresentano scenari futuri ugualmente possibili."
)

doc.add_heading("1.1 Il modello ICON Seamless EPS", level=2)

doc.add_paragraph(
    "Il modello utilizzato in questa analisi e l'ICON Seamless EPS (Ensemble "
    "Prediction System) del servizio meteorologico tedesco (DWD). Questo sistema "
    "produce 40 membri (member 00 = control run + member 01-39 = perturbati). "
    "Ogni membro rappresenta uno scenario equiprobabile dell'evoluzione atmosferica."
)

p = doc.add_paragraph()
run = p.add_run("Principio chiave: ")
run.bold = True
p.add_run(
    "ogni membro ha la stessa probabilita a priori di verificarsi. "
    "Con 40 membri, ciascuno ha peso 1/40 = 2.5%. Questo ci permette "
    "di trattare l'ensemble come un campione da cui stimare distribuzioni "
    "di probabilita."
)

doc.add_heading("1.2 Perche non basta una previsione deterministica?", level=2)

doc.add_paragraph(
    "Una previsione deterministica (singola) fornisce un solo valore: \"domani "
    "la massima sara 10.7 gradi\". Ma non dice nulla sull'incertezza. La massima "
    "potrebbe essere 10.7, ma anche 9.5 o 11.4 con probabilita non trascurabili. "
    "L'ensemble cattura questa incertezza, ed e particolarmente utile per:"
)

items = [
    ("Weather trading", "stimare la probabilita che una soglia venga superata"),
    ("Gestione del rischio", "conoscere gli scenari estremi (code della distribuzione)"),
    ("Decisioni operative", "calibrare azioni in base alla probabilita, non a un singolo numero"),
]
for title_text, desc in items:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(f"{title_text}: ")
    run.bold = True
    p.add_run(desc)


# ═══════════════════════════════════════════════════════════════════════════
# SEZIONE 2 – FORMULE E DEFINIZIONI
# ═══════════════════════════════════════════════════════════════════════════

doc.add_heading("2. Formule e definizioni statistiche", level=1)

# ── Media ────────────────────────────────────────────────────────────────
doc.add_heading("2.1 Media aritmetica", level=2)

doc.add_paragraph(
    "La media dell'ensemble e il valore atteso della temperatura. "
    "Rappresenta la \"migliore stima\" centrale della previsione."
)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("x\u0304 = (1/N) \u00d7 \u2211 x\u1d62     dove N = 40 (numero di membri)")
run.font.size = Pt(12)
run.italic = True

doc.add_paragraph(
    "Esempio: se i 40 valori di temperatura massima del 14 marzo sommano a 427.2, "
    "la media e 427.2 / 40 = 10.68 gradi C."
)

# ── Mediana ───────────────────────────────────────────────────────────────
doc.add_heading("2.2 Mediana", level=2)

doc.add_paragraph(
    "La mediana e il valore centrale quando i dati sono ordinati dal piu piccolo "
    "al piu grande. Con N = 40 (pari), la mediana e la media tra il 20-esimo e "
    "il 21-esimo valore ordinato."
)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Mediana = (x\u2082\u2080 + x\u2082\u2081) / 2     (valori ordinati)")
run.font.size = Pt(12)
run.italic = True

doc.add_paragraph(
    "La mediana e piu robusta della media rispetto ai valori estremi (outlier). "
    "Se media e mediana sono vicine, la distribuzione e approssimativamente simmetrica."
)

# ── Deviazione standard ──────────────────────────────────────────────────
doc.add_heading("2.3 Deviazione standard (spread)", level=2)

doc.add_paragraph(
    "La deviazione standard misura la dispersione dei membri attorno alla media. "
    "Nel contesto ensemble, e chiamata anche \"spread\" e quantifica l'incertezza "
    "della previsione."
)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("s = \u221a[ (1/(N-1)) \u00d7 \u2211 (x\u1d62 - x\u0304)\u00b2 ]")
run.font.size = Pt(12)
run.italic = True

doc.add_paragraph("")
items_std = [
    "Spread basso (es. < 0.5 gradi C): i membri concordano, previsione affidabile",
    "Spread alto (es. > 1.5 gradi C): forte disaccordo, alta incertezza",
    "Lo spread tende ad aumentare con l'orizzonte temporale della previsione",
]
for item in items_std:
    doc.add_paragraph(item, style="List Bullet")

# ── Percentili ───────────────────────────────────────────────────────────
doc.add_heading("2.4 Percentili", level=2)

doc.add_paragraph(
    "I percentili dividono la distribuzione in 100 parti uguali. Il percentile k "
    "e il valore sotto il quale cade il k% dei dati. Sono fondamentali per "
    "descrivere la forma della distribuzione ensemble."
)

doc.add_heading("Definizione formale", level=3)

doc.add_paragraph(
    "Dati N valori ordinati x(1) <= x(2) <= ... <= x(N), il percentile P_k si "
    "calcola come segue:"
)

steps = [
    "1. Calcola l'indice: i = (k / 100) \u00d7 (N - 1) + 1",
    "2. Se i e un intero, P_k = x(i)",
    "3. Se i non e intero, interpola linearmente tra x(\u230ai\u230b) e x(\u2308i\u2309):",
    "   P_k = x(\u230ai\u230b) + (i - \u230ai\u230b) \u00d7 [x(\u2308i\u2309) - x(\u230ai\u230b)]",
]
for step in steps:
    p = doc.add_paragraph(step)
    p.paragraph_format.left_indent = Cm(1)

doc.add_heading("Significato di ogni percentile", level=3)

tbl = doc.add_table(rows=1, cols=3)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl.style = "Table Grid"

hdr = tbl.rows[0]
for i, text in enumerate(["Percentile", "Significato", "Uso nel trading"]):
    hdr.cells[i].text = text
style_header_row(hdr)

perc_data = [
    ("P10 (10-esimo)", "Solo il 10% dei membri prevede un valore\ninferiore a questo. E' il limite basso \"quasi certo\".",
     "Soglia minima conservativa.\nSe il mercato prezza sotto P10, e' probabile\nche sia sottovalutato."),
    ("P25 (25-esimo)\n= 1 quartile", "Il 25% dei membri sta sotto questo valore.\nDelimita il quarto inferiore della distribuzione.",
     "Limite inferiore dello scenario\n\"ragionevolmente probabile\"."),
    ("P50 (50-esimo)\n= Mediana", "Meta' dei membri sta sopra, meta' sotto.\nE' la stima centrale piu' robusta.",
     "Valore di riferimento principale.\nIl mercato dovrebbe convergere qui\nin assenza di bias."),
    ("P75 (75-esimo)\n= 3 quartile", "Solo il 25% dei membri prevede un valore\nsuperiore. Delimita il quarto superiore.",
     "Limite superiore dello scenario\n\"ragionevolmente probabile\"."),
    ("P90 (90-esimo)", "Solo il 10% dei membri prevede un valore\nsuperiore. E' il limite alto \"quasi certo\".",
     "Soglia massima conservativa.\nSe il mercato prezza sopra P90, e' probabile\nche sia sopravvalutato."),
]

for vals in perc_data:
    add_table_row(tbl, vals, center=False)

doc.add_paragraph("")

p = doc.add_paragraph()
run = p.add_run("Intervallo interquartile (IQR) = P75 - P25")
run.bold = True
doc.add_paragraph(
    "L'IQR contiene il 50% centrale dei membri ed e una misura robusta della "
    "dispersione. Nel weather trading, il range P10-P90 (che contiene l'80% dei "
    "membri) e spesso usato come intervallo di confidenza principale."
)

# ── Probabilità empirica ─────────────────────────────────────────────────
doc.add_heading("2.5 Probabilita empirica", level=2)

doc.add_paragraph(
    "La probabilita che la temperatura superi una soglia T* si calcola contando "
    "quanti membri superano quella soglia:"
)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("P(T >= T*) = (numero di membri con T >= T*) / N")
run.font.size = Pt(12)
run.italic = True

doc.add_paragraph(
    "Esempio: se 12 membri su 40 prevedono T max >= 11 gradi C, allora "
    "P(T >= 11) = 12/40 = 30%. Analogamente, per un intervallo:"
)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("P(a <= T < b) = (numero di membri con a <= T < b) / N")
run.font.size = Pt(12)
run.italic = True

# ── KDE ──────────────────────────────────────────────────────────────────
doc.add_heading("2.6 Kernel Density Estimation (KDE)", level=2)

doc.add_paragraph(
    "Con soli 40 valori discreti, la probabilita empirica e \"a gradini\". "
    "Per ottenere una distribuzione continua e liscia, si puo usare la stima a "
    "kernel gaussiano (KDE):"
)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("f(x) = (1/Nh) \u00d7 \u2211 K((x - x\u1d62) / h)")
run.font.size = Pt(12)
run.italic = True

doc.add_paragraph(
    "dove K e il kernel gaussiano e h e la bandwidth (larghezza di banda), che "
    "controlla quanto e \"liscia\" la curva. La bandwidth ottimale per dati "
    "normali e h = 1.06 \u00d7 s \u00d7 N^(-1/5) (regola di Silverman)."
)

doc.add_paragraph(
    "Una volta ottenuta f(x), la probabilita di un intervallo si calcola "
    "integrando: P(a <= T <= b) = integrale da a a b di f(x) dx."
)


# ═══════════════════════════════════════════════════════════════════════════
# SEZIONE 3 – ANALISI DEI DATI
# ═══════════════════════════════════════════════════════════════════════════

doc.add_heading("3. Analisi dei dati: 14-16 Marzo 2026", level=1)

doc.add_paragraph(
    "Di seguito si riportano i risultati per la temperatura massima giornaliera "
    "(temperature_2m_max) prevista dall'ensemble ICON Seamless EPS per Londra "
    "(latitudine 51.505, longitudine 0.055)."
)

date_labels = {
    "2026-03-14": "14 Marzo 2026 (Sabato)",
    "2026-03-15": "15 Marzo 2026 (Domenica)",
    "2026-03-16": "16 Marzo 2026 (Lunedi)",
}

for date_key, label in date_labels.items():
    values = data[date_key]
    stats = compute_stats(values)
    sorted_vals = sorted(values)

    doc.add_heading(f"3.{'123'['2026-03-14,2026-03-15,2026-03-16'.split(',').index(date_key)]}  {label}", level=2)

    # Tabella valori dei 40 membri
    doc.add_heading("Valori dei 40 membri (gradi C)", level=3)

    tbl = doc.add_table(rows=1, cols=8)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"

    hdr = tbl.rows[0]
    for i, text in enumerate(["Member", "T max", "Member", "T max", "Member", "T max", "Member", "T max"]):
        hdr.cells[i].text = text
    style_header_row(hdr)

    for row_idx in range(10):
        row_vals = []
        for col_block in range(4):
            m_idx = col_block * 10 + row_idx
            m_label = f"{m_idx:02d}"
            row_vals.extend([m_label, f"{values[m_idx]:.1f}"])
        add_table_row(tbl, row_vals)

    doc.add_paragraph("")

    # Valori ordinati
    p = doc.add_paragraph()
    run = p.add_run("Valori ordinati: ")
    run.bold = True
    p.add_run(", ".join(f"{v:.1f}" for v in sorted_vals))

    # Tabella statistiche
    doc.add_heading("Statistiche descrittive", level=3)

    tbl2 = doc.add_table(rows=1, cols=2)
    tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl2.style = "Table Grid"

    hdr2 = tbl2.rows[0]
    hdr2.cells[0].text = "Statistica"
    hdr2.cells[1].text = "Valore"
    style_header_row(hdr2)

    stat_rows = [
        ("Media", f"{stats['media']:.2f} C"),
        ("Mediana (P50)", f"{stats['mediana']:.2f} C"),
        ("Deviazione standard", f"{stats['std']:.2f} C"),
        ("Minimo", f"{stats['min']:.1f} C"),
        ("Massimo", f"{stats['max']:.1f} C"),
        ("Range (Max - Min)", f"{stats['range']:.1f} C"),
        ("Percentile 10 (P10)", f"{stats['p10']:.2f} C"),
        ("Percentile 25 (P25 / Q1)", f"{stats['p25']:.2f} C"),
        ("Percentile 75 (P75 / Q3)", f"{stats['p75']:.2f} C"),
        ("Percentile 90 (P90)", f"{stats['p90']:.2f} C"),
        ("IQR (P75 - P25)", f"{stats['p75'] - stats['p25']:.2f} C"),
    ]
    for label, val in stat_rows:
        add_table_row(tbl2, [label, val])

    doc.add_paragraph("")

    # Tabella probabilità
    doc.add_heading("Distribuzione di probabilita", level=3)

    tbl3 = doc.add_table(rows=1, cols=4)
    tbl3.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl3.style = "Table Grid"

    hdr3 = tbl3.rows[0]
    for i, text in enumerate(["Fascia", "N. membri", "Probabilita", "Interpretazione"]):
        hdr3.cells[i].text = text
    style_header_row(hdr3)

    prob_rows = [
        ("T < 10 C", str(stats["n_below_10"]),
         f"{stats['n_below_10']/40*100:.1f}%", "Scenario freddo"),
        ("10 <= T < 11 C", str(stats["n_10_11"]),
         f"{stats['n_10_11']/40*100:.1f}%", "Scenario centrale"),
        ("T >= 11 C", str(stats["n_above_11"]),
         f"{stats['n_above_11']/40*100:.1f}%", "Scenario mite"),
        ("T >= 12 C", str(stats["n_above_12"]),
         f"{stats['n_above_12']/40*100:.1f}%", "Scenario caldo (coda destra)"),
    ]
    for vals in prob_rows:
        add_table_row(tbl3, vals)

    doc.add_paragraph("")


# ═══════════════════════════════════════════════════════════════════════════
# SEZIONE 4 – CONFRONTO TRA I 3 GIORNI
# ═══════════════════════════════════════════════════════════════════════════

doc.add_heading("4. Confronto tra i tre giorni", level=1)

tbl4 = doc.add_table(rows=1, cols=4)
tbl4.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl4.style = "Table Grid"

hdr4 = tbl4.rows[0]
for i, text in enumerate(["Metrica", "14 Marzo", "15 Marzo", "16 Marzo"]):
    hdr4.cells[i].text = text
style_header_row(hdr4)

all_stats = {k: compute_stats(v) for k, v in data.items()}
s14, s15, s16 = all_stats["2026-03-14"], all_stats["2026-03-15"], all_stats["2026-03-16"]

comparison_rows = [
    ("Media", f"{s14['media']:.2f}", f"{s15['media']:.2f}", f"{s16['media']:.2f}"),
    ("Mediana", f"{s14['mediana']:.2f}", f"{s15['mediana']:.2f}", f"{s16['mediana']:.2f}"),
    ("Std Dev (spread)", f"{s14['std']:.2f}", f"{s15['std']:.2f}", f"{s16['std']:.2f}"),
    ("Range", f"{s14['range']:.1f}", f"{s15['range']:.1f}", f"{s16['range']:.1f}"),
    ("P10", f"{s14['p10']:.2f}", f"{s15['p10']:.2f}", f"{s16['p10']:.2f}"),
    ("P90", f"{s14['p90']:.2f}", f"{s15['p90']:.2f}", f"{s16['p90']:.2f}"),
    ("P(T >= 11 C)", f"{s14['n_above_11']/40*100:.1f}%", f"{s15['n_above_11']/40*100:.1f}%", f"{s16['n_above_11']/40*100:.1f}%"),
    ("P(T < 10 C)", f"{s14['n_below_10']/40*100:.1f}%", f"{s15['n_below_10']/40*100:.1f}%", f"{s16['n_below_10']/40*100:.1f}%"),
]

for vals in comparison_rows:
    r = add_table_row(tbl4, vals)
    # prima colonna allineata a sinistra
    r.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

doc.add_paragraph("")

doc.add_heading("4.1 Osservazioni chiave", level=2)

observations = [
    ("Temperatura media stabile: ",
     "la media si mantiene tra 10.68 e 10.87 gradi C nei tre giorni, "
     "suggerendo condizioni simili a livello centrale."),
    ("Incertezza crescente: ",
     "il range aumenta da 1.7 C (14/3) a 2.8 C (15/3) a 3.2 C (16/3). "
     "Questo e tipico: piu ci si allontana nel futuro, piu le traiettorie "
     "dei membri divergono."),
    ("Deviazione standard crescente: ",
     "lo spread passa da ~0.39 C a ~0.53 C a ~0.66 C, confermando "
     "l'aumento progressivo dell'incertezza."),
    ("Distribuzione piu larga il 16/3: ",
     "la probabilita di superare 11 C raddoppia (dal 30% al 50%), e il "
     "rischio di scendere sotto 10 C passa da 0% al 10%. Questo offre "
     "maggiori opportunita di trading ma anche maggiore rischio."),
]

for bold_text, normal_text in observations:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(bold_text)
    run.bold = True
    p.add_run(normal_text)


# ═══════════════════════════════════════════════════════════════════════════
# SEZIONE 5 – COME USARE QUESTI DATI NEL TRADING
# ═══════════════════════════════════════════════════════════════════════════

doc.add_heading("5. Applicazione al weather trading", level=1)

doc.add_heading("5.1 Confronto ensemble vs. prezzo di mercato", level=2)

doc.add_paragraph(
    "Il principio base e confrontare la distribuzione ensemble con il prezzo "
    "quotato dal mercato (ad esempio su CME, Speedwell o piattaforme simili). "
    "Se il mercato quota la temperatura massima del 14 marzo a 11.0 C, e "
    "l'ensemble suggerisce una media di 10.68 C con solo il 30% di probabilita "
    "di superare 11 C, si potrebbe considerare una posizione short (vendita)."
)

doc.add_heading("5.2 Trading su soglie", level=2)

doc.add_paragraph(
    "Molti contratti weather sono binari: pagano se una condizione e soddisfatta "
    "(es. T max > 11 C). L'ensemble fornisce direttamente la probabilita "
    "dell'evento. Se il mercato prezza l'evento al 50% ma l'ensemble dice 30%, "
    "c'e un potenziale edge."
)

doc.add_heading("5.3 Usare lo spread come indicatore di rischio", level=2)

doc.add_paragraph(
    "Uno spread alto indica incertezza: i modelli non concordano. In queste "
    "situazioni si dovrebbe ridurre l'esposizione o usare strategie a basso "
    "rischio. Uno spread basso indica previsione affidabile e giustifica "
    "posizioni piu decise."
)

doc.add_heading("5.4 Multi-modello", level=2)

doc.add_paragraph(
    "L'API Open-Meteo supporta diversi modelli ensemble (ICON-EPS, GFS Ensemble, "
    "ECMWF IFS, ecc.). Confrontare le distribuzioni di modelli diversi aumenta "
    "la robustezza dell'analisi: se tutti concordano, la probabilita e piu "
    "affidabile; se divergono, l'incertezza e reale e va prezzata."
)


# ═══════════════════════════════════════════════════════════════════════════
# SEZIONE 6 – RIFERIMENTO API
# ═══════════════════════════════════════════════════════════════════════════

doc.add_heading("6. Riferimento API Open-Meteo Ensemble", level=1)

doc.add_paragraph("Endpoint utilizzato:")

p = doc.add_paragraph()
run = p.add_run(
    "https://ensemble-api.open-meteo.com/v1/ensemble?"
    "latitude=51.505278&longitude=0.055278"
    "&daily=temperature_2m_max&hourly=temperature_2m"
    "&models=icon_seamless_eps&timezone=auto"
    "&start_date=2026-03-10&end_date=2026-03-18"
)
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x00, 0x66, 0xCC)

doc.add_paragraph("")

doc.add_heading("Parametri principali", level=3)

tbl5 = doc.add_table(rows=1, cols=2)
tbl5.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl5.style = "Table Grid"

hdr5 = tbl5.rows[0]
hdr5.cells[0].text = "Parametro"
hdr5.cells[1].text = "Descrizione"
style_header_row(hdr5)

api_params = [
    ("latitude, longitude", "Coordinate del punto di interesse"),
    ("daily", "Variabili giornaliere (es. temperature_2m_max)"),
    ("hourly", "Variabili orarie (es. temperature_2m)"),
    ("models", "Modello ensemble (icon_seamless_eps, gfs_seamless, ecmwf_ifs025, ...)"),
    ("start_date, end_date", "Intervallo temporale della previsione"),
    ("timezone", "Fuso orario per i timestamp (auto = locale)"),
]
for vals in api_params:
    add_table_row(tbl5, vals, center=False)

doc.add_paragraph("")

doc.add_heading("Struttura della risposta", level=3)

doc.add_paragraph(
    "La risposta JSON contiene un campo \"daily\" (o \"hourly\") con un array "
    "\"time\" e 40 array di valori: \"temperature_2m_max\" (member 00 / control run) "
    "e \"temperature_2m_max_member01\" ... \"temperature_2m_max_member39\" "
    "(i 39 membri perturbati). Ogni array contiene un valore per ogni data/ora "
    "richiesta."
)


# ═══════════════════════════════════════════════════════════════════════════
# SALVATAGGIO
# ═══════════════════════════════════════════════════════════════════════════

output_path = r"c:\Users\Francesco Miserocchi\Documents\Weather trading\Report_Ensemble_Forecast.docx"
doc.save(output_path)
print(f"Report salvato in: {output_path}")
