import pandas as pd
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
forecast_path = os.path.join(BASE_DIR, "Previsioni FM-15 Tutte le Citta 2021-2025 Daily Max.xlsx")
storico_path = os.path.join(BASE_DIR, "Temperature Storiche FM-15 Tutte le Citta 2021-2025 Daily Max.xlsx")
out_path = os.path.join(BASE_DIR, "Report Analisi Previsioni vs Temperature Storiche.docx")

# ── Load data and compute stats for 1GG and 2GG ──
xl_fc = pd.ExcelFile(forecast_path)
xl_st = pd.ExcelFile(storico_path)

results = []
for sheet in xl_fc.sheet_names:
    if sheet not in xl_st.sheet_names:
        continue
    df_fc = pd.read_excel(xl_fc, sheet_name=sheet)
    df_st = pd.read_excel(xl_st, sheet_name=sheet)
    df_fc["Data"] = pd.to_datetime(df_fc["Data"]).dt.date
    df_st["Data"] = pd.to_datetime(df_st["Data"]).dt.date
    m = pd.merge(df_fc, df_st, on="Data", how="inner")

    row = {"city": sheet}

    for label, fc_col_f, fc_col_c in [
        ("1gg", "Max_PrevDay1_F", "Max_PrevDay1_C"),
        ("2gg", "Max_PrevDay2_F", "Max_PrevDay2_C"),
    ]:
        delta_f = (m[fc_col_f] - m["Max_Temperatura_F"]).dropna()
        bias_f = delta_f.mean()
        mae_f = delta_f.abs().mean()
        ratio = abs(bias_f) / mae_f * 100 if mae_f > 0 else 0
        green_f = (delta_f.abs() <= 1).sum() / len(delta_f) * 100
        yellow_f = ((delta_f.abs() > 1) & (delta_f.abs() < 3)).sum() / len(delta_f) * 100
        red_f = (delta_f.abs() >= 3).sum() / len(delta_f) * 100

        delta_c = (m[fc_col_c] - m["Max_Temperatura_C"]).dropna()
        bias_c = delta_c.mean()
        mae_c = delta_c.abs().mean()

        corr_f = delta_f - bias_f
        green_corr = (corr_f.abs() <= 1).sum() / len(corr_f) * 100

        row[f"n_days_{label}"] = len(delta_f)
        row[f"bias_f_{label}"] = round(bias_f, 1)
        row[f"mae_f_{label}"] = round(mae_f, 1)
        row[f"bias_c_{label}"] = round(bias_c, 1)
        row[f"mae_c_{label}"] = round(mae_c, 1)
        row[f"ratio_{label}"] = round(ratio, 0)
        row[f"green_f_{label}"] = round(green_f, 1)
        row[f"yellow_f_{label}"] = round(yellow_f, 1)
        row[f"red_f_{label}"] = round(red_f, 1)
        row[f"green_corr_{label}"] = round(green_corr, 1)
        row[f"improvement_{label}"] = round(green_corr - green_f, 1)

    results.append(row)

df = pd.DataFrame(results)

# ── Helper functions ──
def set_cell_shading(cell, color_hex):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)

def add_styled_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(9)
    return table

# ── Build document ──
doc = Document()

title = doc.add_heading("Report: Analisi Previsioni vs Temperature Storiche", level=0)

doc.add_paragraph(
    "Analisi comparativa tra le temperature massime giornaliere previste (Open-Meteo) "
    "e le temperature massime effettivamente registrate (NOAA ISD, report FM-15) "
    "per 20 citta nel periodo aprile 2021 - agosto 2025. "
    "L'analisi si concentra sulle previsioni emesse 1 giorno prima (PrevDay1) "
    "e 2 giorni prima (PrevDay2) rispetto al giorno di riferimento."
)

# ── 1. Fonti dati ──
doc.add_heading("1. Fonti dei Dati", level=1)

doc.add_heading("Dati storici (temperature registrate)", level=2)
doc.add_paragraph(
    "Le temperature storiche provengono dal database NOAA ISD (Integrated Surface Database). "
    "Sono stati utilizzati esclusivamente i report di tipo FM-15 (METAR), che rappresentano "
    "le osservazioni meteorologiche standard degli aeroporti. I dati originali, in formato UTC, "
    "sono stati convertiti nel fuso orario locale di ciascuna citta utilizzando il modulo Python "
    "zoneinfo con il database IANA (pacchetto tzdata) per gestire correttamente le transizioni "
    "ora legale/solare."
)

doc.add_heading("Dati previsionali (forecast)", level=2)
doc.add_paragraph(
    "Le previsioni meteorologiche provengono dall'API Open-Meteo Previous Runs, che fornisce "
    "lo storico delle previsioni generate dai modelli meteorologici. Le due versioni analizzate sono:"
)
bullet_items = [
    "PrevDay1 (1 giorno prima): la previsione emessa il giorno precedente a quello di riferimento",
    "PrevDay2 (2 giorni prima): la previsione emessa due giorni prima",
]
for item in bullet_items:
    doc.add_paragraph(item, style="List Bullet")

doc.add_paragraph(
    "Il modello utilizzato e 'best_match' di Open-Meteo, con parametro timezone=auto "
    "per ottenere i timestamp in ora locale. Le coordinate utilizzate corrispondono "
    "esattamente alle posizioni degli aeroporti di riferimento."
)

# ── 2. Metodologia ──
doc.add_heading("2. Metodologia", level=1)
doc.add_paragraph(
    "Per ogni citta e ogni giorno, a partire dal 6 aprile 2021, e stata calcolata "
    "la temperatura massima giornaliera sia per i dati storici che per le previsioni. "
    "Il confronto avviene sulla temperatura massima giornaliera in gradi Fahrenheit "
    "(arrotondati all'intero) e Celsius (1 decimale)."
)
doc.add_paragraph("I delta sono calcolati come: Delta = Previsione - Registrata")
doc.add_paragraph(
    "Un delta positivo indica che la previsione sovrastimava la temperatura reale; "
    "un delta negativo indica una sottostima."
)

doc.add_heading("Soglie di accuratezza", level=2)

doc.add_paragraph("Fahrenheit:", style="List Bullet")
p = doc.add_paragraph("", style="List Bullet 2")
p.add_run("Verde").bold = True
p.add_run(": |delta| <= 1 F (previsione accurata)")
p = doc.add_paragraph("", style="List Bullet 2")
p.add_run("Giallo").bold = True
p.add_run(": 1 < |delta| < 3 F (errore moderato)")
p = doc.add_paragraph("", style="List Bullet 2")
p.add_run("Rosso").bold = True
p.add_run(": |delta| >= 3 F (errore significativo)")

doc.add_paragraph("Celsius:", style="List Bullet")
p = doc.add_paragraph("", style="List Bullet 2")
p.add_run("Verde").bold = True
p.add_run(": |delta| <= 1 C")
p = doc.add_paragraph("", style="List Bullet 2")
p.add_run("Giallo").bold = True
p.add_run(": 1 < |delta| < 2 C")
p = doc.add_paragraph("", style="List Bullet 2")
p.add_run("Rosso").bold = True
p.add_run(": |delta| >= 2 C")

doc.add_paragraph(
    "Sono stati esclusi dal conteggio i giorni con valori mancanti nelle colonne delta "
    "(tipicamente concentrati nel 2021, per limitazioni dell'archivio Open-Meteo)."
)

# ── 3. Risultati 1 giorno prima ──
doc.add_heading("3. Risultati: Previsione a 1 Giorno (PrevDay1)", level=1)
doc.add_paragraph(
    "La tabella seguente riporta, per ogni citta, la percentuale di giorni in cui la previsione "
    "emessa 1 giorno prima rientrava nelle soglie di accuratezza in Fahrenheit, "
    "insieme al bias medio e all'errore medio assoluto (MAE)."
)

df_1gg = df.sort_values("green_f_1gg", ascending=False)
headers = ["Citta", "Giorni", "Verde %", "Giallo %", "Rosso %", "Bias F", "MAE F", "Bias C", "MAE C"]
rows = []
for _, r in df_1gg.iterrows():
    rows.append([
        r["city"], r["n_days_1gg"],
        f"{r['green_f_1gg']}%", f"{r['yellow_f_1gg']}%", f"{r['red_f_1gg']}%",
        f"{r['bias_f_1gg']:+.1f}", f"{r['mae_f_1gg']:.1f}",
        f"{r['bias_c_1gg']:+.1f}", f"{r['mae_c_1gg']:.1f}",
    ])
table = add_styled_table(doc, headers, rows)
for i, color in [(2, "92D050"), (3, "FFFF00"), (4, "FF0000")]:
    set_cell_shading(table.rows[0].cells[i], color)

doc.add_paragraph("")
best_1 = df_1gg.iloc[0]
worst_1 = df_1gg.iloc[-1]
doc.add_paragraph(
    f"Previsione a 1 giorno: la citta piu accurata e {best_1['city']} "
    f"({best_1['green_f_1gg']}% verde), la meno accurata e {worst_1['city']} "
    f"({worst_1['green_f_1gg']}% verde)."
)

# ── 4. Risultati 2 giorni prima ──
doc.add_heading("4. Risultati: Previsione a 2 Giorni (PrevDay2)", level=1)
doc.add_paragraph(
    "La tabella seguente riporta gli stessi indicatori per la previsione "
    "emessa 2 giorni prima del giorno di riferimento."
)

df_2gg = df.sort_values("green_f_2gg", ascending=False)
rows = []
for _, r in df_2gg.iterrows():
    rows.append([
        r["city"], r["n_days_2gg"],
        f"{r['green_f_2gg']}%", f"{r['yellow_f_2gg']}%", f"{r['red_f_2gg']}%",
        f"{r['bias_f_2gg']:+.1f}", f"{r['mae_f_2gg']:.1f}",
        f"{r['bias_c_2gg']:+.1f}", f"{r['mae_c_2gg']:.1f}",
    ])
table = add_styled_table(doc, headers, rows)
for i, color in [(2, "92D050"), (3, "FFFF00"), (4, "FF0000")]:
    set_cell_shading(table.rows[0].cells[i], color)

doc.add_paragraph("")
best_2 = df_2gg.iloc[0]
worst_2 = df_2gg.iloc[-1]
doc.add_paragraph(
    f"Previsione a 2 giorni: la citta piu accurata e {best_2['city']} "
    f"({best_2['green_f_2gg']}% verde), la meno accurata e {worst_2['city']} "
    f"({worst_2['green_f_2gg']}% verde)."
)

# ── 5. Confronto 1GG vs 2GG ──
doc.add_heading("5. Confronto: Degradazione da 1 a 2 Giorni", level=1)
doc.add_paragraph(
    "La tabella mostra come la qualita della previsione si degrada passando "
    "da 1 a 2 giorni di anticipo."
)

df_cmp = df.sort_values("green_f_1gg", ascending=False)
headers_cmp = ["Citta", "Verde% 1GG", "Verde% 2GG", "Diff", "MAE 1GG", "MAE 2GG"]
rows_cmp = []
for _, r in df_cmp.iterrows():
    diff = round(r["green_f_2gg"] - r["green_f_1gg"], 1)
    rows_cmp.append([
        r["city"],
        f"{r['green_f_1gg']}%", f"{r['green_f_2gg']}%", f"{diff:+.1f}pp",
        f"{r['mae_f_1gg']:.1f}", f"{r['mae_f_2gg']:.1f}",
    ])
add_styled_table(doc, headers_cmp, rows_cmp)

doc.add_paragraph("")

# ── 6. Diagnosi per 1GG ──
doc.add_heading("6. Diagnosi: Bias Sistematico vs Errore Casuale", level=1)
doc.add_paragraph(
    "Per capire se l'errore previsionale e correggibile, e fondamentale distinguere tra:"
)
p = doc.add_paragraph("", style="List Bullet")
p.add_run("Bias sistematico").bold = True
p.add_run(": la previsione sbaglia sempre nella stessa direzione (es. sempre troppo bassa). "
           "Si corregge con un semplice offset additivo.")
p = doc.add_paragraph("", style="List Bullet")
p.add_run("Errore casuale").bold = True
p.add_run(": la previsione sbaglia a volte in eccesso, a volte in difetto. "
           "Richiede un approccio ensemble (media di piu modelli) per ridurre la varianza.")

doc.add_paragraph(
    "L'indicatore chiave e il rapporto |Bias| / MAE: piu e alto, piu l'errore e sistematico "
    "e correggibile con un offset."
)

# Diagnosi 1GG
doc.add_heading("Previsione a 1 giorno (PrevDay1)", level=2)
headers_d = ["Citta", "Bias F", "MAE F", "|Bias|/MAE", "Diagnosi"]
df_d1 = df.sort_values("ratio_1gg", ascending=False)
rows_d = []
for _, r in df_d1.iterrows():
    ratio = r["ratio_1gg"]
    if ratio >= 70:
        diag = "OFFSET"
    elif ratio >= 40:
        diag = "MISTO"
    else:
        diag = "ENSEMBLE"
    rows_d.append([
        r["city"], f"{r['bias_f_1gg']:+.1f}", f"{r['mae_f_1gg']:.1f}",
        f"{ratio:.0f}%", diag,
    ])
add_styled_table(doc, headers_d, rows_d)
doc.add_paragraph("")

# Diagnosi 2GG
doc.add_heading("Previsione a 2 giorni (PrevDay2)", level=2)
df_d2 = df.sort_values("ratio_2gg", ascending=False)
rows_d2 = []
for _, r in df_d2.iterrows():
    ratio = r["ratio_2gg"]
    if ratio >= 70:
        diag = "OFFSET"
    elif ratio >= 40:
        diag = "MISTO"
    else:
        diag = "ENSEMBLE"
    rows_d2.append([
        r["city"], f"{r['bias_f_2gg']:+.1f}", f"{r['mae_f_2gg']:.1f}",
        f"{ratio:.0f}%", diag,
    ])
add_styled_table(doc, headers_d, rows_d2)
doc.add_paragraph("")

doc.add_paragraph(
    "Le citta classificate come OFFSET hanno un errore prevalentemente unidirezionale. "
    "Le citta ENSEMBLE hanno un bias vicino a zero ma errori distribuiti in entrambe le direzioni. "
    "Le citta MISTE presentano entrambe le componenti."
)

# ── 7. Correzione Offset ──
doc.add_heading("7. Correzione con Offset", level=1)

# 1GG offset
doc.add_heading("Offset su previsione a 1 giorno", level=2)
offset_1gg = df[df["ratio_1gg"] >= 70].sort_values("improvement_1gg", ascending=False)
if len(offset_1gg) > 0:
    doc.add_paragraph(
        f"Per le {len(offset_1gg)} citta con |Bias|/MAE >= 70% a 1 giorno, "
        "e stata simulata una correzione sottraendo il bias medio:"
    )
    headers_o = ["Citta", "Offset F", "Verde% Prima", "Verde% Dopo", "Miglioramento"]
    rows_o = []
    for _, r in offset_1gg.iterrows():
        rows_o.append([
            r["city"], f"{r['bias_f_1gg']:+.1f}",
            f"{r['green_f_1gg']}%", f"{r['green_corr_1gg']}%",
            f"+{r['improvement_1gg']}pp",
        ])
    add_styled_table(doc, headers_o, rows_o)
    doc.add_paragraph("")

# 2GG offset
doc.add_heading("Offset su previsione a 2 giorni", level=2)
offset_2gg = df[df["ratio_2gg"] >= 70].sort_values("improvement_2gg", ascending=False)
if len(offset_2gg) > 0:
    doc.add_paragraph(
        f"Per le {len(offset_2gg)} citta con |Bias|/MAE >= 70% a 2 giorni, "
        "la stessa correzione produce:"
    )
    rows_o2 = []
    for _, r in offset_2gg.iterrows():
        rows_o2.append([
            r["city"], f"{r['bias_f_2gg']:+.1f}",
            f"{r['green_f_2gg']}%", f"{r['green_corr_2gg']}%",
            f"+{r['improvement_2gg']}pp",
        ])
    add_styled_table(doc, headers_o, rows_o2)
    doc.add_paragraph("")

# Seoul note
doc.add_paragraph("")
p = doc.add_paragraph("")
run = p.add_run("Nota su Seoul: ")
run.bold = True
p.add_run(
    "nonostante il bias sia tra i piu alti, il miglioramento con offset e modesto. "
    "Questo indica che, oltre al bias sistematico, Seoul presenta una varianza residua "
    "molto elevata, probabilmente dovuta alla posizione dell'aeroporto di Incheon "
    "(su un'isola costiera con microclima molto diverso dall'entroterra). "
    "Per Seoul, l'offset da solo non basta: servirebbe anche un modello ensemble "
    "o un cambio di stazione di riferimento."
)

# ── 8. Raccomandazioni ──
doc.add_heading("8. Raccomandazioni", level=1)

doc.add_heading("Azioni immediate (offset)", level=2)
doc.add_paragraph("Applicare una correzione offset alle seguenti citta:")

# Combine offset cities from both horizons
all_offset_cities = set(offset_1gg["city"].tolist() + offset_2gg["city"].tolist())
for city in sorted(all_offset_cities):
    r = df[df["city"] == city].iloc[0]
    doc.add_paragraph(
        f"{city}: 1GG {r['bias_f_1gg']:+.1f} F ({r['bias_c_1gg']:+.1f} C), "
        f"2GG {r['bias_f_2gg']:+.1f} F ({r['bias_c_2gg']:+.1f} C)",
        style="List Bullet")

doc.add_heading("Azioni consigliate (ensemble)", level=2)
doc.add_paragraph(
    "Per le citta con errore prevalentemente casuale (ENSEMBLE e MISTO), "
    "valutare l'utilizzo di modelli ensemble dall'API Open-Meteo. "
    "Un ensemble (media o mediana di piu modelli) riduce la varianza "
    "dell'errore previsionale, migliorando l'accuratezza soprattutto "
    "dove il bias e gia vicino a zero."
)
doc.add_paragraph("Citta prioritarie per l'ensemble (ordinate per verde% 1GG crescente):")
ensemble_cities = df[df["ratio_1gg"] < 70].sort_values("green_f_1gg")
for _, r in ensemble_cities.iterrows():
    doc.add_paragraph(
        f"{r['city']} (1GG: {r['green_f_1gg']}% verde, MAE: {r['mae_f_1gg']} F; "
        f"2GG: {r['green_f_2gg']}% verde, MAE: {r['mae_f_2gg']} F)",
        style="List Bullet")

doc.add_heading("Considerazioni generali", level=2)
doc.add_paragraph(
    "Il confronto e basato su temperature massime giornaliere derivate da dati orari. "
    "I report METAR FM-15 sono osservazioni puntuali (tipicamente ogni 1-3 ore) e potrebbero "
    "non catturare il vero picco giornaliero se questo si verifica tra due osservazioni. "
    "Questo introduce un errore strutturale che penalizza le citta con poche osservazioni "
    "giornaliere o con picchi di temperatura molto brevi."
)

# ── Save ──
doc.save(out_path)
print(f"Report salvato: {out_path}")
